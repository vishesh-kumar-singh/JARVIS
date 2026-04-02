"""
Google Drive Service Wrapper.
Provides tools to search for files, download PDFs/Docs, extract their text, 
and run RAG (Retrieval-Augmented Generation) queries directly on downloaded documents.
"""
import sys
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from langchain.schema import Document
from RAG import RAG
import io
import os


import docx
import PyPDF2

import re

def clean_text(text):
    # Remove multiple consecutive newlines
    text = re.sub(r'\n\s*\n', '\n', text)
    # Replace line breaks not after punctuation with space
    text = re.sub(r'(?<![.!?])\n', ' ', text)
    # Collapse multiple spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


class GoogleDrive:
    def __init__(self, credentials):
        try:
            self.service = build("drive", "v3", credentials=credentials)
        except:
            self.service = None
            pass # "Unable to make connection with Drive")

    def search_files(self, keywords=None, max_results=10):
        if not self.service:
            return []

        if not keywords:
            return []

        # Build the Drive query with AND conditions
        query_parts = [f"fullText contains '{kw}'" for kw in keywords]
        query = " or ".join(query_parts)
        # Exclude folders
        query = f"({query}) and mimeType != 'application/vnd.google-apps.folder'"

        results = self.service.files().list(
            q=query,
            pageSize=max_results,
            fields="files(id, name, mimeType, modifiedTime, owners)"
        ).execute()
    
        return results.get("files", [])

    def download_file(self, file_id, filepath="Temporary/downloaded_file"):
        try:
            os.makedirs(os.path.dirname(filepath), exist_ok=True)

            file = self.service.files().get(fileId=file_id).execute()
            mime_type = file.get("mimeType")

            if mime_type.startswith("application/vnd.google-apps"):
                export_mime = None
                if mime_type == "application/vnd.google-apps.document":
                    export_mime = "application/pdf"
                    filepath += ".pdf"
                elif mime_type == "application/vnd.google-apps.spreadsheet":
                    export_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    filepath += ".xlsx"
                elif mime_type == "application/vnd.google-apps.presentation":
                    export_mime = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
                    filepath += ".pptx"

                if not export_mime:
                    pass # f"Export not supported for mime type {mime_type}")
                    return None

                request = self.service.files().export_media(fileId=file_id, mimeType=export_mime)
            else:
                # guess extension if possible
                if mime_type == "application/pdf":
                    filepath += ".pdf"
                elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    filepath += ".docx"
                elif mime_type == "text/plain":
                    filepath += ".txt"
                request = self.service.files().get_media(fileId=file_id)

            fh = io.FileIO(filepath, "wb")
            downloader = MediaIoBaseDownload(fh, request)

            done = False
            while not done:
                status, done = downloader.next_chunk()

            return filepath

        except HttpError as error:
            pass # f"An error occurred: {error}")
            return None

    def rag_on_file(self, filepaths, query):
        """Download a file, extract text, and run RAG on it."""
        pass # "Fecting resources from Drive")
        docs=[]
        text=""
        for filepath in filepaths:
            if filepath.endswith(".pdf"):
                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += clean_text(page.extract_text()) or ""
            elif filepath.endswith(".docx"):
                doc = docx.Document(filepath)
                for para in doc.paragraphs:
                    text += clean_text(para.text) + "\n"
            elif filepath.endswith(".txt"):
                with open(filepath, "r", encoding="utf-8") as f:
                    text = clean_text(f.read())
            else:
                return f"Text extraction not supported for {filepath}"
            docs.append(Document(page_content=text, metadata={"file": filepath}))

        return RAG(docs, query)

    def get_results(self, query, keywords, max_results=5):
        files = self.search_files(keywords, max_results)
        if not files:
            return "No files found."
        
        filepaths=[]
        for i,file in enumerate(files):
            pass # f"Found file: {file['name']} (ID: {file['id']})")
            path=self.download_file(file['id'], f"Temporary/downloaded_file{i}")
            if path:
                filepaths.append(path)
        
        if not filepaths:
            return "Failed to download any relevant files for searching."
        
        context = self.rag_on_file(filepaths, query)

        return context
        